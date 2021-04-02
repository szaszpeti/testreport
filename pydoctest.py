import docx 
from docx.shared import Inches
import pprint
from docx import Document
import os
from copy import deepcopy
import mammoth
import codecs
from docxtpl import DocxTemplate
import bs4

#file location
"""
SET THIS MANUALY BEFORE RUNING THE CODE

name the offshore report files:
    M35_A_Internal.docx
    M35_A_External.docx
    M35_B_EInternal.doxc
    ......
    
"""


driver = "F"
turbine = "M77"
blade = "A"
inspection = "external"
#tableInspectionNumber = 7

 

offshoreReport = (""+ driver + ":\\" + turbine + "\\" + turbine + "_" + blade + "_" + inspection + ".docx")
if os.path.isfile(""+ driver + ":\\" + turbine + "\\" + "Rotor blades inspection report " + inspection  + " " + turbine + ".docx"):
    reportLayout = (""+ driver + ":\\" + turbine + "\\" + "Rotor blades inspection report " + inspection  + " " + turbine + ".docx")
    reportLayoutSavePath = (""+ driver + ":\\" + turbine + "\\" + "Rotor blades inspection report " + inspection  + " " + turbine + ".docx")
else:   
    reportLayout = (""+ driver + ":\\" + 'Vorlage' + "\\" + "Rotor blades inspection report " + inspection + " layout.docx")
    reportLayoutSavePath = (""+ driver + ":\\" + turbine + "\\" + "Rotor blades inspection report " + inspection  + " " + turbine + ".docx")
findingLayoutExternal = (""+ driver + ":\\" + 'Vorlage' + "\\" + "FindingLayoutExternal.docx")
findingLayoutInternal = (""+ driver + ":\\" + 'Vorlage' + "\\" + "FindingLayoutInternal2.docx")
findingLayoutSavePath = (""+ driver + ":\\" + turbine + "\\" + "mylayout" + "_" + turbine + "_" + blade + "_" + inspection + ".docx")

htmlOffshoreReport = (""+ driver + ":\\" + turbine + "\\" + turbine + "_" + blade + "_" + inspection + ".html")

"""
OPENING THE REPORT
"""

docOffshore = Document(offshoreReport)


"""SITE_DETAILS table"""
cell
tableSiteDetails = docOffshore.tables[0]

site_number = tableSiteDetails.cell(2,1).text
turbine_serial_number = tableSiteDetails.cell(3,1).text


blade_position_number = tableSiteDetails.cell(6,1).text
blade_serial_number = tableSiteDetails.cell(7,1).text

print("BLADE NUMBER: " + blade_position_number + "_" + blade_serial_number)

#goes to layout 908767/M34
wtg_number = turbine_serial_number + "/" + site_number

print("BLADE SERIAL NUMBER: " + wtg_number)


"""
---------Create Folder for The Images 
Blade_A_53612000_External

"""
imageDirectory = (""+ driver + ":\\" + turbine + "\\" + "Blade" + "_" + blade + "_" + blade_serial_number + "_" + inspection.capitalize())
if not os.path.isdir(""+ driver + ":\\" + turbine + "\\" + "Blade" + "_" + blade + "_" + blade_serial_number + "_" + inspection.capitalize()):
    os.mkdir(""+ driver + ":\\" + turbine + "\\" + "Blade" + "_" + blade + "_" + blade_serial_number + "_" + inspection.capitalize())

#----------------------------------------------------

"""WORK_TEAM table"""

tableWorkTeamDate = docOffshore.tables[1]

working_team = tableWorkTeamDate.cell(0,1).text
date = tableWorkTeamDate.cell(3,1).text

print("WORKING TEAM: " + working_team)
print("DATE OF INSPECTION: " + date)
print("INSPECTION TYPE: " + inspection)

"""----------------------------------------------------"""

"""EXTERNAL INSPECTION RESULT"""
"""
1. external report => tables[7]
   
2. external from MIXREPORT report => tables[8]
   
3. internal all reports => tables [7]

"""

tablesOffshore = docOffshore.tables

def find_damage_table(table):
    for x in range(len(table)):
        for row in table[x].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text == "Distance from LE (TE =0%)" and inspection == "external":
                        return table[x]
                        break
                    elif paragraph.text == "Comments" and inspection == "internal":
                        return table[x]
                    break
    
print ("INSPECTION RESULT TABLE: " + str(tablesOffshore.index((find_damage_table(tablesOffshore)))))

tableInspectionResult = docOffshore.tables[tablesOffshore.index((find_damage_table(tablesOffshore)))]


#put all the cell values to list "data"
data=[]

for row in tableInspectionResult.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            data.append(paragraph.text)


        

#separete the key and value (paros pozicio key, paratlan value)
new = [data[i::2] for i in range(2)]

for x in new[1]:
    print ("This is elemt : " + x)

valueList = new[1]
damageListProcentComments = []

    

if inspection == "external":
   for x in range(int(len(valueList)/13)):
       damageListProcentComments.append(valueList[0:9])
       del valueList[:13]
    
else:
     for x in range(int(len(valueList)/12)):
       damageListProcentComments.append(valueList[0:8])
       del valueList[:12]



for x in range(len(damageListProcentComments)):
    if damageListProcentComments[x][-1] == "":
        damageListProcentComments[x][-1] = "comments n.a."
        



print(" ")
print("FINDINGS: " + str(len(damageListProcentComments)))
pprint.pprint(damageListProcentComments)


#list for percent values only
if inspection == "external":
    procent_list = []
    for list in damageListProcentComments:
        procent_list.append(list[7])


commentsOfInspector = []

if inspection == "external":
    for list in damageListProcentComments:
        commentsOfInspector.append(list[8])
        
else:
    for list in damageListProcentComments:
        commentsOfInspector.append(list[7])


for x in commentsOfInspector:
    print("This is comments of inspector : " + x)
    


"""
 ---- getting the LPS table conttent
"""

lpsValues = []

if inspection == "external":
    tableLPS = docOffshore.tables[4]
    for row in tableLPS.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                try:
                    lpsValues.append(format(float(paragraph.text), '.3f'))
                except ValueError:
                    pass
for x in lpsValues:
    print("LPS VLUES: " + str(x))




"""
OPENING THE LAYOUT----------------------------------------------------
"""

#open the layout document
layout = Document(reportLayout)

#Overview WTG (table 1)
tableOverviewWtgData = layout.tables[0]

wtg_number_layout_cell = tableOverviewWtgData.cell(3,2)

wtg_number_layout_cell.text = wtg_number


#Overview Blade Data (table2)
tableOverviewBladeData = layout.tables[1]
set_number = tableOverviewBladeData.cell(2,1)
blade_number_A = tableOverviewBladeData.cell(3,1)
blade_number_B = tableOverviewBladeData.cell(4,1)
blade_number_C = tableOverviewBladeData.cell(5,1)
inspection_period = tableOverviewBladeData.cell(6,1)
name_of_technicians_I = tableOverviewBladeData.cell(8,1)
name_of_technicians_II = tableOverviewBladeData.cell(8,1)

if blade_position_number == "A":
    blade_number_A.text = blade_serial_number
elif blade_position_number == "B":
     blade_number_B.text = blade_serial_number
else:
    blade_number_C.text = blade_serial_number
        
inspection_period.text = date
name_of_technicians_I.text = working_team


#Get
if blade == "A":
    tableOverviewInspectionSummary = layout.tables[4]
elif blade == "B":
    tableOverviewInspectionSummary = layout.tables[5]
else:
    tableOverviewInspectionSummary = layout.tables[6]

listOfDamages = []
#prepare composite list for Overview List
for d in damageListProcentComments:
    del d[7:]
    if inspection == "external":
        d.insert(len(d),"external")
    else:
        d.insert(len(d),"internal")
    d.insert(0, blade_serial_number)
    listOfDamages.append(d)

pprint.pprint(listOfDamages)


"""FILL DATA IN OVERVIEW LIST"""

row = 1
cell = 1
for x in range(len(listOfDamages)):
    for i in listOfDamages[x]:
        tableOverviewInspectionSummary.cell(x+1, listOfDamages[x].index(i)).text = i.replace("mm","")
        cell += 1

#get rid of M
for x in range(len(listOfDamages)):
    tableOverviewInspectionSummary.cell(x+1, 2).text = tableOverviewInspectionSummary.cell(x+1, 2).text.replace("M", "")
        
    
      #add damage category on surface laminate etc...       
for x in range(len(listOfDamages)):
    if tableOverviewInspectionSummary.cell(x+1, 5). text == "1":
        tableOverviewInspectionSummary.cell(x+1, 4).text =  tableOverviewInspectionSummary.cell(x+1, 4).text + category_1
    elif tableOverviewInspectionSummary.cell(x+1, 5). text == "2":
        tableOverviewInspectionSummary.cell(x+1, 4).text =  tableOverviewInspectionSummary.cell(x+1, 4).text + category_2
    elif tableOverviewInspectionSummary.cell(x+1, 5). text == "3":
        tableOverviewInspectionSummary.cell(x+1, 4).text =  tableOverviewInspectionSummary.cell(x+1, 4).text + category_3
    else:
        pass          



"""
CREATE INSPECTION RESULT IN DIFFERENT DOCX EXTERNAL
"""
#Get Inspection Result Table

'''
findign layout test, copy the findig table and past it as many times as many
damage we founmd 
'''

#open finding layout   
if inspection == "external":
    finding_layout_external = Document(findingLayoutExternal)
    finding_1 = finding_layout_external.tables[0]
    
    
    #copy paste finding table as many damages found
    for x in range(int(len(listOfDamages))):
        tbl, p= finding_1._tbl, finding_layout_external.paragraphs[x]._p
        finding_layout_external.add_paragraph('')
        #finding_layout.paragraphs[x].add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        new_tbl = deepcopy(tbl)
        p.addnext(new_tbl)
        
    
    
    #put data into finding tables
    for x in range(len(listOfDamages)):
            #findign number
            finding_layout_external.tables[x].cell(1,0).text = str(listOfDamages[x][0] + '-' + str(x+1))
            #position SS/TE/PS/LE
            finding_layout_external.tables[x].cell(0,2).text = str(listOfDamages[x][3]) 
            #category
            finding_layout_external.tables[x].cell(1,2).text = str(listOfDamages[x][5])                                
            #distance from to in M
            finding_layout_external.tables[x].cell(0,4).text = str(listOfDamages[x][2]).replace("M","")                               
            finding_layout_external.tables[x].cell(1,4).text = str(listOfDamages[x][2]).replace("M","")                                
            #procent
            finding_layout_external.tables[x].cell(0,6).text = str(procent_list[x])
                                    
            finding_layout_external.tables[x].cell(1,6).text = str(procent_list[x])  
            #size                           
            finding_layout_external.tables[x].cell(0,8).text = str(listOfDamages[x][6].replace("mm","") + 'x' + listOfDamages[x][7].replace("mm",""))
            #damage type
            if finding_layout_external.tables[x].cell(1,2).text == "1":                        
                finding_layout_external.tables[x].cell(2,3).text = str(listOfDamages[x][4] + ";" + " on surface coating")
            if finding_layout_external.tables[x].cell(1,2).text == "2":                        
                finding_layout_external.tables[x].cell(2,3).text = str(listOfDamages[x][4] + ";" + " on surface laminate")
            if finding_layout_external.tables[x].cell(1,2).text == "3":                        
                finding_layout_external.tables[x].cell(2,3).text = str(listOfDamages[x][4] + ";" + " on multilayer laminate")
            #comments of Inspector
            finding_layout_external.tables[x].cell(5,1).text = str(commentsOfInspector[x])
    
    
    
    finding_layout_external.save(findingLayoutSavePath)
    
if inspection == "internal":
    finding_layout_internal = Document(findingLayoutInternal)
    finding_1 = finding_layout_internal.tables[0]
    
    
    #copy paste finding table as many damages found
    for x in range(int(len(listOfDamages))):
        tbl, p= finding_1._tbl, finding_layout_internal.paragraphs[x]._p
        finding_layout_internal.add_paragraph('')
        #finding_layout.paragraphs[x].add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        new_tbl = deepcopy(tbl)
        p.addnext(new_tbl)
        
    
    
    #put data into finding tables
    for x in range(len(listOfDamages)):
            #findign number
            finding_layout_internal.tables[x].cell(1,0).text = str(listOfDamages[x][0] + '-' + str(x+1))
            #position SS/TE/PS/LE
            finding_layout_internal.tables[x].cell(0,2).text = str(listOfDamages[x][3]) 
            #category
            finding_layout_internal.tables[x].cell(1,2).text = str(listOfDamages[x][5])                                
            #distance from to in M
            finding_layout_internal.tables[x].cell(0,5).text = str(listOfDamages[x][2]).replace("M","")                               
            finding_layout_internal.tables[x].cell(1,5).text = str(listOfDamages[x][2]).replace("M","")                                
            #procent
            #damage type                            
            finding_layout_internal.tables[x].cell(0,10).text = str(listOfDamages[x][6].replace("mm",""))
            finding_layout_internal.tables[x].cell(1,10).text = str(listOfDamages[x][7].replace("mm",""))
            if finding_layout_internal.tables[x].cell(1,3).text == "1":                        
                finding_layout_internal.tables[x].cell(2,3).text = str(listOfDamages[x][4] + ";" + " on surface coating")
            if finding_layout_internal.tables[x].cell(1,3).text == "2":                        
                finding_layout_internal.tables[x].cell(2,3).text = str(listOfDamages[x][4] + ";" + " on surface laminate")
            if finding_layout_internal.tables[x].cell(1,3).text == "3":                        
                finding_layout_internal.tables[x].cell(2,3).text = str(listOfDamages[x][4] + ";" + " on multilayer laminate")
            #comments of Inspector
        
            finding_layout_internal.tables[x].cell(5,2).text = "Comment of inspector: " + str(commentsOfInspector[x])
            
            #pictures

            p0 = finding_layout_internal.tables[x].cell(4,0).add_paragraph()
            p1 = finding_layout_internal.tables[x].cell(4,6).add_paragraph()
            p2 = finding_layout_internal.tables[x].cell(4,10).add_paragraph()
            
            pic1 = p0.add_run()
            pic2 = p1.add_run()
            pic3 = p2.add_run()
            
            if os.path.isfile(F""+ driver + ":\\" + turbine + "\\Blade" + "_" + blade_position_number + "_" + blade_serial_number + "_" + inspection + "\\" + blade_serial_number + "_" + str(x+1) + "_" + "1.jpeg"):
                pic1.add_picture(F""+ driver + ":\\" + turbine + "\\Blade" + "_" + blade_position_number + "_" + blade_serial_number + "_" + inspection + "\\" + blade_serial_number + "_" + str(x+1) + "_" + "1.jpeg", width=Inches(1.5))
            if os.path.isfile(F""+ driver + ":\\" + turbine + "\\Blade" + "_" + blade_position_number + "_" + blade_serial_number + "_" + inspection + "\\" + blade_serial_number + "_" + str(x+1) + "_" + "2.jpeg"):
                pic2.add_picture(F""+ driver + ":\\" + turbine + "\\Blade" + "_" + blade_position_number + "_" + blade_serial_number + "_" + inspection + "\\" + blade_serial_number + "_" + str(x+1) + "_" + "2.jpeg", width=Inches(1.5))
            if os.path.isfile(F""+ driver + ":\\" + turbine + "\\Blade" + "_" + blade_position_number + "_" + blade_serial_number + "_" + inspection + "\\" + blade_serial_number + "_" + str(x+1) + "_" + "3.jpeg"):
                pic3.add_picture(F""+ driver + ":\\" + turbine + "\\Blade" + "_" + blade_position_number + "_" + blade_serial_number + "_" + inspection + "\\" + blade_serial_number + "_" + str(x+1) + "_" + "3.jpeg", width=Inches(1.5))
#
# =============================================================================
#            
# =============================================================================
 
            
             
# =============================================================================
            for x in range(len(listOfDamages)):
                t = finding_layout_internal.tables[x]
                tbl, p  = t._tbl, layout.paragraphs[x]._p   
                layout.add_paragraph('')
                new_tbl = deepcopy(tbl)
                


            
        

            
    finding_layout_internal.save(findingLayoutSavePath)



    
    #layoutInspectionResult.save("F:\python\WPy64-3740\documents\layout_inspection_result.docx")
layout.save(reportLayoutSavePath)





templateDoc = DocxTemplate(reportLayoutSavePath)


context = { 'tableA' : new_tbl }


templateDoc.render(context)
templateDoc.save(reportLayoutSavePath)


"""

------------------------------ RENAME PICTURES -----------------------

"""
# =============================================================================
# from docx import Document
# from docx.shared import Inches
# 
# doc = Document(findingLayoutSavePath)
# tables = doc.tables
# p = tables[0].rows[4].cells[0].add_paragraph()
# r = p.add_run()
# r.add_picture('F:\\M30\\18.jpeg',width=Inches(1.0), height=Inches(1.7))
# p = tables[1].rows[4].cells[0].add_paragraph()
# r = p.add_run()
# r.add_picture('F:\\M30\\18.jpeg',width=Inches(4.0), height=Inches(.7))
# doc.save('F:\\M30\\addImage.docx')
# 
# =============================================================================


# =============================================================================
# 
# import os
# os.getcwd()
# collection = "F:\M35\Blade_A_xxxxxx_External"
# for i, filename in enumerate(os.listdir(collection)):
#     os.rename("F:\M35\Blade_A_xxxxxx_External\\" + filename, blade_serial_number + "_" + str(i+1) + ".jpg")
# =============================================================================



# =============================================================================
# 

 
# =============================================================================
# b= open (htmlOffshoreReport, "wb")
# with open(offshoreReport, "rb") as docx_file:
#     result = mammoth.convert_to_html(docx_file)
#     html = result.value # The generated HTML
#     messages = result.messages # Any messages, such as warnings during conversion
# 
# b.write(html.encode('utf8'))
# b.close()
# 
# =============================================================================
# 
# =============================================================================
import os
import shutil





class ImageWriter(object):
    def __init__(self, output_dir):
        self._output_dir = output_dir
        self._image_number = 1

    def __call__(self, element):
        extension = element.content_type.partition("/")[2]
        image_filename = "{0}.{1}".format(self._image_number, extension)
        with open(os.path.join(self._output_dir, image_filename), "wb") as image_dest:
            with element.open() as image_source:
                shutil.copyfileobj(image_source, image_dest)

        self._image_number += 1

        return {"src": image_filename}


convert_image = mammoth.images.inline(ImageWriter(imageDirectory))
result = mammoth.convert_to_html(offshoreReport,convert_image=convert_image)




# =============================================================================
# with codecs.open(htmlOffshoreReport, encoding="utf-8", mode="w") as text_file:
#     text_file.write(result.value)
# 
# =============================================================================


# =============================================================================
# with open(htmlOffshoreReport, 'r') as f:
#   webpage = f.read().decode('utf-8')
#   
# soup = bs4.BeautifulSoup(webpage, "html.parser")
# 
# 
# for a in soup.find_all('a'):
#     if a.img:
#         print(a.img['src'])
# 
# =============================================================================




"""
MAMMOTH
"""
#mammoth "f:\python\WPy64-3740\documents\m30\m30external_report.docx" "f:\python\WPy64-3740\documents\m30\m30external_report.html"




#mammoth document.docx --output-dir=output-dir

    








#https://github.com/python-openxml/python-docx/blob/master/features/steps/paragraph.py
#https://grokonez.com/python/how-to-read-write-word-docx-files-in-python-docx-module

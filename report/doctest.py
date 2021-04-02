from docx import Document
from docx.shared import Inches



def report(operation, inspection):
	doc = Document('external.docx')
	table = doc.tables[1]
	print(table.cell(1,1).text)	
	table.cell(2,1).text = operation.set_nummber
	table.cell(3,1).text = operation.blade_number
	
	table.cell(8,1).text = operation.name_1
	table.cell(9,1).text = operation.name_2

	damage = doc.tables[2]
	print(len(inspection))




	doc.save('external2.docx')